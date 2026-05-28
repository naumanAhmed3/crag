"""One-shot generator for the binary parts of the sample corpus.

Writes `corpus/handbook.docx` and `corpus/benchmarks.xlsx`. Re-run any
time the content below changes. Idempotent.
"""

from __future__ import annotations

from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent / "corpus"


# ── handbook.docx ─────────────────────────────────────────────────────


HANDBOOK_TITLE = "Vantage Holdings — Knowledge Search Handbook"

HANDBOOK_SECTIONS: list[tuple[str, str]] = [
    (
        "1. Scope",
        "This handbook documents the day-to-day operating contract for "
        "the offline knowledge-search service. It is the source of truth "
        "for what the service is allowed to do, what it must refuse, and "
        "how operators recover from common failure modes.",
    ),
    (
        "2. Source-grounding requirement",
        "Every answer the service returns must cite at least one passage "
        "from the indexed corpus. Where no passage supports a claim, the "
        "service must reply with the verbatim string \"I don't have enough "
        'information in the indexed corpus to answer this." Operators '
        "monitor the rate of this refusal as the primary trust metric.",
    ),
    (
        "3. Retention and deletion",
        "Documents are retained in the index for as long as they remain "
        "on the authoritative file share. Removal from the share triggers "
        "automatic eviction on the next ingest cycle, with the corresponding "
        "chunk identifiers logged to the audit trail. For urgent removal "
        "ahead of an ingest cycle, operators use scripts/eviction.py with "
        "an explicit reason string.",
    ),
    (
        "4. Audit trail",
        "Every ingest, eviction, and incident response is logged to a "
        "structured JSONL file under $CRAG_DATA_DIR/audit/. Audit log "
        "integrity is verified by an external job that hashes each day's "
        "rotated file and checks the hash against a controlled register.",
    ),
    (
        "5. Approved models",
        "Only models listed in the approved-model register may be loaded. "
        "The current approved set is bge-small-en-v1.5 for embedding, "
        "bge-reranker-v2-m3 for reranking, and qwen2.5:7b-instruct-q4_K_M "
        "for generation. Substitutions require change-management approval "
        "and a re-run of the full evaluation gold set.",
    ),
    (
        "6. Off-corpus questions",
        "Where a user asks a question the corpus cannot answer, the "
        "service replies with the refusal string described in section 2. "
        "Operators must not configure the service to answer from the "
        "model's parametric knowledge under any circumstance.",
    ),
]


def build_handbook() -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    title = doc.add_heading(HANDBOOK_TITLE, level=0)
    title.alignment = 1  # centred
    doc.add_paragraph(
        "Internal — operational reference. Distributed only to operators "
        "responsible for the knowledge-search service."
    )

    for heading, body in HANDBOOK_SECTIONS:
        doc.add_heading(heading, level=1)
        para = doc.add_paragraph(body)
        for run in para.runs:
            run.font.size = Pt(11)

    out = ROOT / "handbook.docx"
    doc.save(str(out))
    return out


# ── benchmarks.xlsx ───────────────────────────────────────────────────


BENCH_ROWS: list[tuple[str, str, float, float, float]] = [
    # (configuration, hardware, recall@5, mrr, p95_latency_ms)
    ("dense only, BGE-small", "M3 Pro, CPU", 0.620, 0.471, 410.0),
    ("dense + BM25 + RRF", "M3 Pro, CPU", 0.745, 0.598, 480.0),
    ("dense + BM25 + RRF + rerank", "M3 Pro, CPU", 0.840, 0.712, 1240.0),
    ("dense + BM25 + RRF + rerank", "RTX 3060", 0.840, 0.712, 380.0),
]


def build_benchmarks() -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    wb = Workbook()
    ws = wb.active
    ws.title = "Retrieval bake-off"

    header = ["Configuration", "Hardware", "Recall@5", "MRR", "p95 latency (ms)"]
    ws.append(header)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="left")

    for row in BENCH_ROWS:
        ws.append(list(row))

    for col_letter, width in (("A", 38), ("B", 16), ("C", 12), ("D", 10), ("E", 18)):
        ws.column_dimensions[col_letter].width = width

    # Second sheet: index-size baselines.
    ws2 = wb.create_sheet("Index-size baselines")
    ws2.append(
        ["Corpus size (docs)", "Chunks", "Disk (MB)", "RAM at load (MB)", "Cold-ingest minutes"]
    )
    for cell in ws2[1]:
        cell.font = Font(bold=True)
    for row in [
        (1_000, 3_800, 18.4, 412.0, 0.6),
        (10_000, 38_000, 168.0, 480.0, 4.7),
        (100_000, 380_000, 1_640.0, 1_120.0, 41.0),
        (1_000_000, 3_800_000, 16_200.0, 9_800.0, 410.0),
    ]:
        ws2.append(list(row))
    for col_letter, width in (("A", 22), ("B", 14), ("C", 14), ("D", 20), ("E", 22)):
        ws2.column_dimensions[col_letter].width = width

    out = ROOT / "benchmarks.xlsx"
    wb.save(str(out))
    return out


def main() -> None:
    h = build_handbook()
    b = build_benchmarks()
    print(f"wrote {h}")
    print(f"wrote {b}")


if __name__ == "__main__":
    main()
