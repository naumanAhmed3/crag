"""Format-aware document parsing.

Each parser returns a list of `RawText` blocks. A block is a contiguous
piece of readable text with optional locator metadata (page number, sheet
name, heading path) — the locator becomes part of the citation surfaced
to the user.

Format dispatch is by file extension. Adding a new format means writing
a function with signature `(path: Path) -> list[RawText]` and registering
it in `PARSERS`. No heavy dependencies (no `unstructured.io`, no Java);
all wheels are pure-Python or pip-installable on POSIX without system
packages.
"""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class RawText:
    """A contiguous piece of text from a source file."""

    text: str
    locator: str  # human-readable position, e.g. "p. 4" or "Sheet1!A1:Z31"


# ── PDF ─────────────────────────────────────────────────────────────────


def _parse_pdf(path: Path) -> list[RawText]:
    import pymupdf  # local import keeps cli startup fast

    blocks: list[RawText] = []
    with pymupdf.open(path) as doc:
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                blocks.append(RawText(text=text, locator=f"p. {i}"))
    return blocks


# ── DOCX ────────────────────────────────────────────────────────────────


def _parse_docx(path: Path) -> list[RawText]:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if line:
                parts.append(line)
    body = "\n\n".join(parts).strip()
    return [RawText(text=body, locator="document")] if body else []


# ── XLSX ────────────────────────────────────────────────────────────────


def _parse_xlsx(path: Path) -> list[RawText]:
    from openpyxl import load_workbook

    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    blocks: list[RawText] = []
    for sheet in wb.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(RawText(text="\n".join(rows), locator=f"sheet: {sheet.title}"))
    wb.close()
    return blocks


# ── HTML ────────────────────────────────────────────────────────────────


def _parse_html(path: Path) -> list[RawText]:
    from selectolax.parser import HTMLParser

    raw = path.read_text(encoding="utf-8", errors="ignore")
    tree = HTMLParser(raw)
    for tag in ("script", "style", "nav", "footer"):
        for node in tree.css(tag):
            node.decompose()
    body = tree.body
    text = (body.text(separator="\n", strip=True) if body else tree.text(strip=True)).strip()
    return [RawText(text=text, locator="document")] if text else []


# ── Markdown / plain text / source code ─────────────────────────────────


def _parse_text(path: Path) -> list[RawText]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    return [RawText(text=text, locator="document")] if text else []


# ── Dispatch ────────────────────────────────────────────────────────────

PARSERS: dict[str, Callable[[Path], list[RawText]]] = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".xlsx": _parse_xlsx,
    ".html": _parse_html,
    ".htm": _parse_html,
    ".md": _parse_text,
    ".markdown": _parse_text,
    ".txt": _parse_text,
    ".rst": _parse_text,
    ".py": _parse_text,
    ".js": _parse_text,
    ".ts": _parse_text,
    ".json": _parse_text,
    ".yaml": _parse_text,
    ".yml": _parse_text,
    ".sql": _parse_text,
    ".sh": _parse_text,
    ".go": _parse_text,
    ".rs": _parse_text,
}

SUPPORTED_EXTENSIONS = frozenset(PARSERS)


def parse(path: Path) -> list[RawText]:
    """Parse a single file. Raises KeyError if the extension is unsupported."""
    parser = PARSERS.get(path.suffix.lower())
    if parser is None:
        raise KeyError(f"No parser registered for extension {path.suffix!r}")
    return parser(path)
